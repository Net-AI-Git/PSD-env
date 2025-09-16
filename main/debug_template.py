from docx import Document
import os

TEMPLATE_PATH = 'Data/format.docx'

def inspect_document(filepath):
    """
    Loads a .docx file and prints the text and underlying XML of each paragraph.

    The 'Why': This script is a diagnostic tool to understand the internal structure
    of a .docx template file. It helps debug issues where the main script fails
    to find expected elements like placeholders for titles, images, or captions.

    The 'What': It loads a specified .docx document and iterates through every
    paragraph. For each paragraph, it prints both the user-visible text and the
    raw underlying XML that the python-docx library interacts with. This allows
    developers to see the exact tags and structures they need to target in code.

    Args:
        filepath (str): The path to the .docx file to inspect.
    """
    if not os.path.exists(filepath):
        print(f"ERROR: File not found at '{filepath}'")
        return

    print(f"--- Inspecting Document: {filepath} ---")
    try:
        doc = Document(filepath)
        print(f"Successfully loaded document. Found {len(doc.paragraphs)} paragraphs.\\n")

        for i, p in enumerate(doc.paragraphs):
            print(f"--- Paragraph {i} ---")
            print(f"Text: '{p.text}'")
            print(f"XML: {p._p.xml}\\n")
            
        print("--- Inspection Complete ---")

    except Exception as e:
        print(f"An error occurred while reading the document: {e}")


if __name__ == "__main__":
    inspect_document(TEMPLATE_PATH)
