import os
import copy
import traceback
from typing import List, Tuple
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _add_image_block(doc: Document, image_path: str, caption_template: Paragraph, num_spacers: int):
    """
    Adds a complete 'image block' (image, caption, and spacing) to the document.

    The 'Why': This function encapsulates the full logic for a single image entry,
    mirroring the structure found in the user's template. It handles the image,
    the complex caption replacement, and the visual spacing, ensuring consistency.

    The 'What': It performs a sequence of stable operations:
    1. Adds the picture.
    2. Deep copies and appends the caption template paragraph.
    3. Finds and replaces the 'A title' placeholder in the new caption. If not
       found, it safely appends the title to the end.
    4. Adds a specific number of empty paragraphs for spacing.

    Args:
        doc (Document): The main Word document object.
        image_path (str): The full path to the image file to add.
        caption_template (Paragraph): The template paragraph for the caption.
        num_spacers (int): The number of empty paragraphs to add after the caption.
    """
    try:
        # 1. Add the picture
        doc.add_picture(image_path, width=Inches(6.0))

        # 2. Add a placeholder paragraph to ensure correct positioning
        placeholder_p = doc.add_paragraph()

        # 3. Create a deep copy of the caption template's XML element
        caption_element = copy.deepcopy(caption_template._p)
        
        # 4. Replace the placeholder paragraph's element with the caption element
        placeholder_p._p.getparent().replace(placeholder_p._p, caption_element)
        
        # 5. Get a reference to the new caption (which is now the last paragraph)
        #    and modify its text
        new_caption_paragraph = doc.paragraphs[-1]
        image_title = os.path.splitext(os.path.basename(image_path))[0]
        
        placeholder_found = False
        for run in new_caption_paragraph.runs:
            if 'A title' in run.text:
                run.text = run.text.replace('A title', image_title)
                placeholder_found = True
                break
        
        if not placeholder_found:
            new_caption_paragraph.add_run(f" {image_title}")

        # 6. Add spacer paragraphs
        for _ in range(num_spacers):
            doc.add_paragraph()

    except Exception:
        print(f"\\n{'='*80}")
        print(f"FATAL ERROR: Could not process and add image block for: {image_path}")
        traceback.print_exc()
        print(f"{'='*80}\\n")

def _analyze_template(template_path: str) -> Tuple[Paragraph, int]:
    """
    Analyzes the template file to extract the caption paragraph and spacer count.
    """
    try:
        template_doc = Document(template_path)
    except Exception as e:
        print(f"FATAL ERROR: Could not load the template file at '{template_path}'. Error: {e}")
        return None, 0

    caption_template = None
    caption_index = -1
    caption_field_identifier = '<w:instrText>איור</w:instrText>'

    for i, p in enumerate(template_doc.paragraphs):
        if caption_field_identifier in p._p.xml:
            caption_template = p
            caption_index = i
            break
    
    if caption_template is None:
        return None, 0

    num_spacers = 0
    for i in range(caption_index + 1, len(template_doc.paragraphs)):
        if template_doc.paragraphs[i].text.strip() == "":
            num_spacers += 1
        else:
            break
            
    return caption_template, num_spacers

def create_images_document(image_paths: List[str], output_path: str, document_name: str = "Image_Document.docx") -> None:
    """
    Creates a Word document from a template, following a precise, robust plan.

    The 'Why': This function is the master controller for creating the Word report.
    It is designed for reliability by first analyzing a template, validating its
    contents, and then building a new document from scratch based on that analysis.

    The 'What':
    1. Analyzes 'Data/format.docx' to find the caption template and count empty
       spacer paragraphs that follow it. Fails with a clear error if not found.
    2. Creates a new, blank document.
    3. Sorts images into 'main' and 'details' lists.
    4. Loops through the sorted lists, calling a helper function to add a full
       'image block' (picture, caption, spacers) for each image.
    5. Saves the final document.

    Args:
        image_paths (List[str]): Full paths to the image files.
        output_path (str): The directory where the Word document will be saved.
        document_name (str, optional): The filename for the output Word document.
    """
    template_path = 'Data/format.docx'

    # 1. Analyze the template
    caption_template, num_spacers = _analyze_template(template_path)
    
    if caption_template is None:
        print(f"FATAL ERROR: Could not find a valid caption template in '{template_path}'. Process stopped.")
        return

    # 2. Create new document
    doc = Document()

    # 3. Sort and add image blocks
    details_images = sorted([p for p in image_paths if 'details' in os.path.basename(p).lower()])
    main_images = sorted([p for p in image_paths if 'details' not in os.path.basename(p).lower()])

    print(f"Processing {len(main_images)} main images...")
    for image_path in main_images:
        _add_image_block(doc, image_path, caption_template, num_spacers)

    print(f"Processing {len(details_images)} details images...")
    for image_path in details_images:
        _add_image_block(doc, image_path, caption_template, num_spacers)

    # --- 4. Apply Final Document-Wide Formatting ---
    print("Applying final formatting (Center Align, RTL)...")
    for paragraph in doc.paragraphs:
        # 1. Set paragraph-level alignment and direction
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.rtl = True

        # 2. Set run-level direction to enforce RTL on the actual text
        for run in paragraph.runs:
            run.font.rtl = True

    # --- 5. Save the Final Document ---
    save_path = os.path.join(output_path, document_name)
    try:
        doc.save(save_path)
        print(f"Document successfully saved to {save_path}")
    except Exception as e:
        print(f"FATAL ERROR: Could not save the final document to '{save_path}'.")
        traceback.print_exc()
